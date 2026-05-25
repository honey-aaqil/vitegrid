from __future__ import annotations

import os
import re
import statistics
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pymupdf
from docling.document_converter import DocumentConverter
from docx import Document as DocxDocument
from docx.shared import RGBColor

try:
    from rapidocr_onnxruntime import RapidOCR  # type: ignore[import-not-found]
except ImportError:
    RapidOCR = None  # type: ignore[assignment,misc]


@dataclass
class ParsedDocument:
    markdown: str
    tables: list[dict[str, Any]] = field(default_factory=list)
    images: list[dict[str, Any]] = field(default_factory=list)
    page_count: int = 0
    raw: dict[str, Any] = field(default_factory=dict)


@dataclass
class TextSpan:
    text: str
    bbox: tuple[float, float, float, float]
    font: str
    size_pt: float
    color_hex: str
    bold: bool
    italic: bool
    page_index: int

    def to_dict(self) -> dict[str, Any]:
        return {
            "text": self.text,
            "bbox": list(self.bbox),
            "font": self.font,
            "size_pt": self.size_pt,
            "color_hex": self.color_hex,
            "bold": self.bold,
            "italic": self.italic,
            "page_index": self.page_index,
        }


@dataclass
class PageLayout:
    page_index: int
    page_width_pt: float
    page_height_pt: float
    spans: list[TextSpan]
    image_bytes: bytes


@dataclass
class PdfExtraction:
    pages: list[PageLayout]
    is_scanned: bool


_ocr_engine: Any = None


def _get_ocr_engine() -> Any:
    global _ocr_engine
    if _ocr_engine is None:
        if RapidOCR is None:
            raise RuntimeError(
                "rapidocr-onnxruntime is not installed. Run: pip install rapidocr-onnxruntime"
            )
        _ocr_engine = RapidOCR()
    return _ocr_engine


def run_ocr_on_image(image_bytes: bytes, image_dpi: int = 200) -> list[TextSpan]:
    engine = _get_ocr_engine()
    import io
    from PIL import Image as PILImage
    import numpy as np

    img = PILImage.open(io.BytesIO(image_bytes)).convert("RGB")
    arr = np.array(img)
    raw = engine(arr)
    result = raw[0] if isinstance(raw, tuple) else raw
    if not result:
        return []
    pt_per_px = 72.0 / image_dpi
    spans: list[TextSpan] = []
    for item in result:
        if not item:
            continue
        box = item[0] if len(item) > 0 else None
        text = item[1] if len(item) > 1 else None
        score_raw = item[2] if len(item) > 2 else 1.0
        try:
            score = float(score_raw)
        except (TypeError, ValueError):
            score = 1.0
        if not box or not text or not str(text).strip():
            continue
        if score < 0.4:
            continue
        text = str(text)
        xs = [p[0] for p in box]
        ys = [p[1] for p in box]
        x0_px, x1_px = min(xs), max(xs)
        y0_px, y1_px = min(ys), max(ys)
        height_px = max(1.0, y1_px - y0_px)
        spans.append(
            TextSpan(
                text=text.strip(),
                bbox=(
                    x0_px * pt_per_px,
                    y0_px * pt_per_px,
                    x1_px * pt_per_px,
                    y1_px * pt_per_px,
                ),
                font="OCR",
                size_pt=height_px * pt_per_px * 0.75,
                color_hex="#1a1a1a",
                bold=False,
                italic=False,
                page_index=0,
            )
        )
    return spans


def _bbox_iou(a: tuple[float, float, float, float], b: tuple[float, float, float, float]) -> float:
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    ix0 = max(ax0, bx0)
    iy0 = max(ay0, by0)
    ix1 = min(ax1, bx1)
    iy1 = min(ay1, by1)
    if ix1 <= ix0 or iy1 <= iy0:
        return 0.0
    inter = (ix1 - ix0) * (iy1 - iy0)
    area_a = max(1.0, (ax1 - ax0) * (ay1 - ay0))
    area_b = max(1.0, (bx1 - bx0) * (by1 - by0))
    return inter / min(area_a, area_b)


def merge_native_and_ocr(native: list[TextSpan], ocr: list[TextSpan], iou_threshold: float = 0.5) -> list[TextSpan]:
    if not ocr:
        return list(native)
    if not native:
        return list(ocr)
    merged: list[TextSpan] = list(native)
    for o in ocr:
        overlap = False
        for n in native:
            if _bbox_iou(o.bbox, n.bbox) >= iou_threshold:
                overlap = True
                break
        if not overlap:
            merged.append(o)
    return merged


def augment_extraction_with_ocr(extraction: PdfExtraction) -> PdfExtraction:
    for page in extraction.pages:
        if extraction.is_scanned or len(page.spans) < 10:
            ocr_spans = run_ocr_on_image(page.image_bytes)
            for s in ocr_spans:
                s.page_index = page.page_index
            page.spans = merge_native_and_ocr(page.spans, ocr_spans)
    return extraction


_converter: DocumentConverter | None = None


def _get_converter() -> DocumentConverter:
    global _converter
    if _converter is None:
        _converter = DocumentConverter()
    return _converter


def parse_document(file_path: str | Path) -> ParsedDocument:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Source file not found: {path}")

    result = _get_converter().convert(str(path))
    doc = result.document

    markdown = doc.export_to_markdown()
    raw_dict = doc.export_to_dict()

    tables: list[dict[str, Any]] = []
    for table in getattr(doc, "tables", []) or []:
        cells: list[list[str]] = []
        try:
            for row in table.data.grid:
                cells.append([cell.text for cell in row])
        except AttributeError:
            cells = []
        tables.append({
            "rows": len(cells),
            "cols": len(cells[0]) if cells else 0,
            "cells": cells,
        })

    images: list[dict[str, Any]] = []
    for pic in getattr(doc, "pictures", []) or []:
        bbox = getattr(pic, "bbox", None)
        images.append({
            "id": getattr(pic, "self_ref", None),
            "bbox": {
                "x": getattr(bbox, "l", 0) if bbox else 0,
                "y": getattr(bbox, "t", 0) if bbox else 0,
                "width": (getattr(bbox, "r", 0) - getattr(bbox, "l", 0)) if bbox else 0,
                "height": (getattr(bbox, "b", 0) - getattr(bbox, "t", 0)) if bbox else 0,
            },
        })

    return ParsedDocument(
        markdown=markdown,
        tables=tables,
        images=images,
        page_count=len(getattr(doc, "pages", []) or []),
        raw=raw_dict,
    )


def pdf_pages_to_images(pdf_path: str | Path, dpi: int = 200, max_pages: int = 10) -> list[bytes]:
    path = Path(pdf_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF not found: {path}")
    doc = pymupdf.open(str(path))
    images: list[bytes] = []
    zoom = dpi / 72.0
    matrix = pymupdf.Matrix(zoom, zoom)
    for page in doc[:max_pages]:
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        images.append(pix.tobytes("png"))
    doc.close()
    return images


def _int_color_to_hex(value: int | None) -> str:
    if value is None:
        return "#1a1a1a"
    r = (value >> 16) & 0xFF
    g = (value >> 8) & 0xFF
    b = value & 0xFF
    return f"#{r:02x}{g:02x}{b:02x}"


def _font_flags(flags: int) -> tuple[bool, bool]:
    bold = bool(flags & 16)
    italic = bool(flags & 2)
    return bold, italic


def extract_pdf_layout(
    pdf_path: str | Path,
    image_dpi: int = 200,
    max_pages: int = 10,
) -> PdfExtraction:
    path = Path(pdf_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF not found: {path}")
    doc = pymupdf.open(str(path))
    pages: list[PageLayout] = []
    total_chars = 0
    zoom = image_dpi / 72.0
    matrix = pymupdf.Matrix(zoom, zoom)
    for idx, page in enumerate(doc[:max_pages]):
        text_dict = page.get_text("dict")
        spans: list[TextSpan] = []
        for block in text_dict.get("blocks", []):
            if block.get("type", 0) != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = (span.get("text") or "").strip()
                    if not text:
                        continue
                    total_chars += len(text)
                    bold, italic = _font_flags(int(span.get("flags", 0)))
                    bbox = tuple(span.get("bbox", (0, 0, 0, 0)))
                    spans.append(
                        TextSpan(
                            text=text,
                            bbox=(bbox[0], bbox[1], bbox[2], bbox[3]),
                            font=str(span.get("font", "")),
                            size_pt=float(span.get("size", 0)),
                            color_hex=_int_color_to_hex(span.get("color")),
                            bold=bold,
                            italic=italic,
                            page_index=idx,
                        )
                    )
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        pages.append(
            PageLayout(
                page_index=idx,
                page_width_pt=float(page.rect.width),
                page_height_pt=float(page.rect.height),
                spans=spans,
                image_bytes=pix.tobytes("png"),
            )
        )
    doc.close()
    is_scanned = total_chars < 20
    return PdfExtraction(pages=pages, is_scanned=is_scanned)


@dataclass
class DocxBlock:
    type: str
    text: str | None
    items: list[str] | None
    rows: list[list[str]] | None
    font: str | None
    size_pt: float | None
    color_hex: str | None
    bold: bool
    italic: bool
    align: str | None


def _docx_color(color: RGBColor | None) -> str | None:
    if color is None:
        return None
    try:
        return f"#{color:06x}" if isinstance(color, int) else f"#{int(color):06x}"
    except (TypeError, ValueError):
        return None


def _docx_alignment(value: Any) -> str | None:
    if value is None:
        return None
    name = str(value).split(".")[-1].lower()
    if name in ("left", "center", "centre", "right", "justify", "both"):
        if name in ("center", "centre"):
            return "center"
        if name == "both":
            return "justify"
        return name
    return None


def extract_docx_layout(docx_path: str | Path) -> list[DocxBlock]:
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")
    doc = DocxDocument(str(path))
    blocks: list[DocxBlock] = []
    pending_list: list[str] = []
    pending_list_meta: dict[str, Any] | None = None

    def flush_list() -> None:
        nonlocal pending_list_meta
        if pending_list and pending_list_meta is not None:
            blocks.append(
                DocxBlock(
                    type="list",
                    text=None,
                    items=list(pending_list),
                    rows=None,
                    font=pending_list_meta.get("font"),
                    size_pt=pending_list_meta.get("size_pt"),
                    color_hex=pending_list_meta.get("color_hex"),
                    bold=pending_list_meta.get("bold", False),
                    italic=pending_list_meta.get("italic", False),
                    align=pending_list_meta.get("align"),
                )
            )
        pending_list.clear()
        pending_list_meta = None

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = (para.style.name or "").lower() if para.style else ""
        first_run = next((r for r in para.runs if r.text.strip()), None)
        font_name = first_run.font.name if first_run and first_run.font.name else None
        size_pt = first_run.font.size.pt if first_run and first_run.font.size else None
        color_hex = _docx_color(first_run.font.color.rgb) if first_run and first_run.font.color else None
        bold = bool(first_run.bold) if first_run and first_run.bold is not None else "bold" in style_name
        italic = bool(first_run.italic) if first_run and first_run.italic is not None else False
        align = _docx_alignment(para.alignment)

        is_list_item = "list" in style_name or "bullet" in style_name
        if is_list_item and text:
            pending_list.append(text)
            if pending_list_meta is None:
                pending_list_meta = {
                    "font": font_name,
                    "size_pt": size_pt,
                    "color_hex": color_hex,
                    "bold": bold,
                    "italic": italic,
                    "align": align,
                }
            continue

        flush_list()
        if not text:
            continue

        if "heading" in style_name or "title" in style_name:
            blocks.append(
                DocxBlock(
                    type="heading",
                    text=text,
                    items=None,
                    rows=None,
                    font=font_name,
                    size_pt=size_pt or 18.0,
                    color_hex=color_hex,
                    bold=True if bold is None else bold,
                    italic=italic,
                    align=align,
                )
            )
        else:
            blocks.append(
                DocxBlock(
                    type="paragraph",
                    text=text,
                    items=None,
                    rows=None,
                    font=font_name,
                    size_pt=size_pt,
                    color_hex=color_hex,
                    bold=bold,
                    italic=italic,
                    align=align,
                )
            )

    flush_list()

    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            blocks.append(
                DocxBlock(
                    type="table",
                    text=None,
                    items=None,
                    rows=rows,
                    font=None,
                    size_pt=10.0,
                    color_hex=None,
                    bold=False,
                    italic=False,
                    align=None,
                )
            )

    return blocks


@dataclass
class ClassifiedBlock:
    block_type: str
    text: str | None
    items: list[str] | None
    rows: list[list[str]] | None
    font: str | None
    size_pt: float | None
    color_hex: str | None
    bold: bool
    italic: bool
    align: str
    bbox: tuple[float, float, float, float]
    page_index: int


_BULLET_CHARS = "•‣◦⁃∙·▪▫■□—-*–"
_BULLET_RE = re.compile(rf"^[{re.escape(_BULLET_CHARS)}]\s+")
_NUMBER_RE = re.compile(r"^(\d{1,3}[.)]|\([a-zA-Z0-9]{1,3}\)|[a-zA-Z][.)])\s+")


def _line_align(line_x0: float, line_x1: float, page_width: float, tol: float = 8.0) -> str:
    left_gap = line_x0
    right_gap = page_width - line_x1
    if abs(left_gap - right_gap) <= tol and left_gap > 30:
        return "center"
    if right_gap < tol and left_gap > 100:
        return "right"
    return "left"


def _group_lines(spans: list[TextSpan]) -> list[list[TextSpan]]:
    if not spans:
        return []
    sorted_spans = sorted(spans, key=lambda s: (round(s.bbox[1], 1), s.bbox[0]))
    lines: list[list[TextSpan]] = []
    for span in sorted_spans:
        height = max(span.size_pt, 6.0)
        placed = False
        cy = (span.bbox[1] + span.bbox[3]) / 2
        for line in lines:
            ref = line[0]
            ref_cy = (ref.bbox[1] + ref.bbox[3]) / 2
            if abs(cy - ref_cy) <= height * 0.5:
                line.append(span)
                placed = True
                break
        if not placed:
            lines.append([span])
    for line in lines:
        line.sort(key=lambda s: s.bbox[0])
    lines.sort(key=lambda line: (line[0].bbox[1] + line[0].bbox[3]) / 2)
    return lines


def _line_text(line: list[TextSpan]) -> str:
    if not line:
        return ""
    parts: list[str] = []
    prev_x1: float | None = None
    for span in line:
        if prev_x1 is not None and span.bbox[0] - prev_x1 > span.size_pt * 0.2:
            parts.append(" ")
        parts.append(span.text)
        prev_x1 = span.bbox[2]
    return "".join(parts).strip()


def _line_bbox(line: list[TextSpan]) -> tuple[float, float, float, float]:
    x0 = min(s.bbox[0] for s in line)
    y0 = min(s.bbox[1] for s in line)
    x1 = max(s.bbox[2] for s in line)
    y1 = max(s.bbox[3] for s in line)
    return x0, y0, x1, y1


def _group_blocks(lines: list[list[TextSpan]]) -> list[list[list[TextSpan]]]:
    if not lines:
        return []
    blocks: list[list[list[TextSpan]]] = [[lines[0]]]
    for prev, current in zip(lines, lines[1:]):
        prev_bbox = _line_bbox(prev)
        cur_bbox = _line_bbox(current)
        prev_size = max(s.size_pt for s in prev)
        cur_size = max(s.size_pt for s in current)
        smaller_size = min(prev_size, cur_size)
        line_height = smaller_size * 1.2
        gap = cur_bbox[1] - prev_bbox[3]
        size_ratio = max(prev_size, cur_size) / max(smaller_size, 1.0)
        prev_text = _line_text(prev)
        cur_text = _line_text(current)
        prev_bold = sum(1 for s in prev if s.bold) > len(prev) / 2
        cur_bold = sum(1 for s in current if s.bold) > len(current) / 2

        if size_ratio > 1.25:
            blocks.append([current])
        elif prev_bold != cur_bold and gap > line_height * 0.3:
            blocks.append([current])
        elif gap > line_height * 0.8:
            blocks.append([current])
        elif _looks_like_list_marker(prev_text) != _looks_like_list_marker(cur_text):
            blocks.append([current])
        else:
            blocks[-1].append(current)
    return blocks


def _looks_like_list_marker(text: str) -> bool:
    return bool(_BULLET_RE.match(text)) or bool(_NUMBER_RE.match(text))


def _strip_list_marker(text: str) -> str:
    m = _BULLET_RE.match(text)
    if m:
        return text[m.end():].strip()
    m = _NUMBER_RE.match(text)
    if m:
        return text[m.end():].strip()
    return text.strip()


def _detect_table(block_lines: list[list[TextSpan]]) -> list[list[str]] | None:
    if len(block_lines) < 2:
        return None
    line_col_starts: list[list[float]] = []
    for line in block_lines:
        starts: list[float] = []
        for i, span in enumerate(line):
            if i == 0:
                starts.append(span.bbox[0])
                continue
            prev = line[i - 1]
            if span.bbox[0] - prev.bbox[2] > prev.size_pt * 1.2:
                starts.append(span.bbox[0])
        line_col_starts.append(starts)
    multi_col_lines = [s for s in line_col_starts if len(s) >= 2]
    if len(multi_col_lines) < 2:
        return None
    target = max(len(s) for s in multi_col_lines)
    column_anchors = next(s for s in multi_col_lines if len(s) == target)
    rows: list[list[str]] = []
    for line, starts in zip(block_lines, line_col_starts):
        if len(starts) < 2:
            continue
        cells: list[list[str]] = [[] for _ in column_anchors]
        for span in line:
            idx = 0
            for j, anchor in enumerate(column_anchors):
                if span.bbox[0] >= anchor - 4:
                    idx = j
            cells[idx].append(span.text)
        rows.append([" ".join(parts).strip() for parts in cells])
    if len(rows) < 2:
        return None
    return rows


def _majority_font(spans: list[TextSpan]) -> str | None:
    fonts = [s.font for s in spans if s.font]
    if not fonts:
        return None
    return max(set(fonts), key=fonts.count)


def _majority_color(spans: list[TextSpan]) -> str | None:
    colors = [s.color_hex for s in spans if s.color_hex]
    if not colors:
        return None
    return max(set(colors), key=colors.count)


def _avg_size(spans: list[TextSpan]) -> float:
    sizes = [s.size_pt for s in spans if s.size_pt]
    return statistics.mean(sizes) if sizes else 11.0


def _is_bold(spans: list[TextSpan]) -> bool:
    if not spans:
        return False
    bold_count = sum(1 for s in spans if s.bold)
    return bold_count > len(spans) / 2


def _classify_block_spans(
    block_lines: list[list[TextSpan]],
    page_width_pt: float,
    page_index: int,
    body_size: float,
    heading_threshold: float,
) -> ClassifiedBlock | None:
    block_spans = [s for line in block_lines for s in line]
    if not block_spans:
        return None
    bbox = _line_bbox(block_spans)
    font = _majority_font(block_spans)
    color = _majority_color(block_spans)
    size_pt = round(_avg_size(block_spans), 1)
    bold = _is_bold(block_spans)
    italic = sum(1 for s in block_spans if s.italic) > len(block_spans) / 2
    line_texts = [_line_text(line) for line in block_lines]
    joined = " ".join(line_texts).strip()
    x0, _, x1, _ = bbox
    align = _line_align(x0, x1, page_width_pt)

    list_lines = [t for t in line_texts if _looks_like_list_marker(t)]
    if len(list_lines) >= 2 and len(list_lines) / len(line_texts) >= 0.7:
        items = [_strip_list_marker(t) for t in line_texts if t]
        return ClassifiedBlock(
            block_type="list",
            text=None,
            items=items,
            rows=None,
            font=font,
            size_pt=size_pt,
            color_hex=color,
            bold=bold,
            italic=italic,
            align=align,
            bbox=bbox,
            page_index=page_index,
        )

    table_rows = _detect_table(block_lines)
    if table_rows is not None:
        return ClassifiedBlock(
            block_type="table",
            text=None,
            items=None,
            rows=table_rows,
            font=font,
            size_pt=size_pt,
            color_hex=color,
            bold=bold,
            italic=italic,
            align=align,
            bbox=bbox,
            page_index=page_index,
        )

    is_heading = (
        len(block_lines) == 1
        and len(joined) < 120
        and (size_pt >= heading_threshold or (bold and size_pt >= body_size))
    )
    return ClassifiedBlock(
        block_type="heading" if is_heading else "paragraph",
        text=joined,
        items=None,
        rows=None,
        font=font,
        size_pt=size_pt,
        color_hex=color,
        bold=bold,
        italic=italic,
        align=align,
        bbox=bbox,
        page_index=page_index,
    )


def _classify_pdf_layout_v1(extraction: PdfExtraction) -> list[ClassifiedBlock]:
    classified: list[ClassifiedBlock] = []
    all_spans = [s for page in extraction.pages for s in page.spans]
    if not all_spans:
        return classified
    body_size = statistics.median([s.size_pt for s in all_spans if s.size_pt])
    heading_threshold = body_size * 1.15
    for page in extraction.pages:
        lines = _group_lines(page.spans)
        block_lines = _group_blocks(lines)
        for block in block_lines:
            cb = _classify_block_spans(
                block, page.page_width_pt, page.page_index, body_size, heading_threshold
            )
            if cb is not None:
                classified.append(cb)
    return classified


# ----- Parser v2 (projection columns + cross-row table anchors) ----------------


def detect_columns(
    spans: list[TextSpan],
    page_width_pt: float,
    *,
    min_gutter_pt: float = 18.0,
    min_column_pt: float = 60.0,
) -> list[tuple[float, float]]:
    """Find column x-ranges via horizontal projection of span coverage.

    Sweep-line over text x-intervals: any contiguous gutter of >= `min_gutter_pt`
    that has zero text coverage separates two columns. Columns thinner than
    `min_column_pt` are absorbed into their nearest neighbor to avoid splitting
    on a single floating word.
    """
    if not spans:
        return [(0.0, page_width_pt)]

    events: list[tuple[float, int]] = []
    for s in spans:
        x0, _, x1, _ = s.bbox
        if x1 <= x0:
            continue
        events.append((float(x0), +1))
        events.append((float(x1), -1))
    events.sort()

    coverage_runs: list[tuple[float, float]] = []
    active = 0
    run_start: float | None = None
    for x, delta in events:
        prev = active
        active += delta
        if prev == 0 and active > 0:
            run_start = x
        elif prev > 0 and active == 0 and run_start is not None:
            coverage_runs.append((run_start, x))
            run_start = None
    if run_start is not None:
        coverage_runs.append((run_start, float(page_width_pt)))

    # Merge runs separated by gutters smaller than min_gutter_pt.
    merged: list[tuple[float, float]] = []
    for start, end in coverage_runs:
        if merged and start - merged[-1][1] < min_gutter_pt:
            merged[-1] = (merged[-1][0], end)
        else:
            merged.append((start, end))

    # Drop columns thinner than min_column_pt by merging into nearest neighbor.
    if len(merged) > 1:
        i = 0
        while i < len(merged):
            start, end = merged[i]
            if end - start < min_column_pt:
                if i == 0:
                    merged[1] = (start, merged[1][1])
                    merged.pop(0)
                elif i == len(merged) - 1:
                    merged[i - 1] = (merged[i - 1][0], end)
                    merged.pop(i)
                    i -= 1
                else:
                    # Merge into whichever neighbor has the smaller gutter.
                    left_gap = start - merged[i - 1][1]
                    right_gap = merged[i + 1][0] - end
                    if left_gap <= right_gap:
                        merged[i - 1] = (merged[i - 1][0], end)
                    else:
                        merged[i + 1] = (start, merged[i + 1][1])
                    merged.pop(i)
                    continue
            i += 1

    return merged if merged else [(0.0, page_width_pt)]


def _spans_in_column(spans: list[TextSpan], col: tuple[float, float]) -> list[TextSpan]:
    cx0, cx1 = col
    return [
        s
        for s in spans
        if (s.bbox[0] + s.bbox[2]) / 2.0 >= cx0 - 0.5
        and (s.bbox[0] + s.bbox[2]) / 2.0 <= cx1 + 0.5
    ]


def _detect_table_v2(block_lines: list[list[TextSpan]]) -> list[list[str]] | None:
    """Detect unbordered tables by cross-row x-anchor alignment.

    Strategy: cluster the x-start of every span across all lines. If we can find
    >= 2 anchor clusters that each contain >= 2 lines' worth of spans, and the
    block has >= 2 lines whose spans align to >= 2 of those anchors, that block
    is a table. This catches modern unbordered tables where v1 fails because it
    requires intra-line whitespace gaps.
    """
    if len(block_lines) < 2:
        return None
    anchors: list[float] = []
    eps = 8.0
    for line in block_lines:
        for span in line:
            placed = False
            for i, a in enumerate(anchors):
                if abs(span.bbox[0] - a) <= eps:
                    anchors[i] = (a + span.bbox[0]) / 2.0
                    placed = True
                    break
            if not placed:
                anchors.append(span.bbox[0])
    anchors.sort()

    line_anchor_hits: list[list[int]] = []
    for line in block_lines:
        hits: list[int] = []
        for span in line:
            for i, a in enumerate(anchors):
                if abs(span.bbox[0] - a) <= eps:
                    if i not in hits:
                        hits.append(i)
                    break
        line_anchor_hits.append(sorted(hits))

    multi = [h for h in line_anchor_hits if len(h) >= 2]
    if len(multi) < 2:
        return None
    target = sorted({i for h in multi for i in h})
    if len(target) < 2:
        return None

    rows: list[list[str]] = []
    for line, hits in zip(block_lines, line_anchor_hits):
        if len(hits) < 2:
            continue
        cells: list[list[str]] = [[] for _ in target]
        anchor_xs = [anchors[i] for i in target]
        for span in line:
            j = 0
            for k, ax in enumerate(anchor_xs):
                if span.bbox[0] >= ax - eps:
                    j = k
            cells[j].append(span.text)
        rows.append([" ".join(parts).strip() for parts in cells])
    return rows if len(rows) >= 2 else None


def _classify_pdf_layout_v2(extraction: PdfExtraction) -> list[ClassifiedBlock]:
    classified: list[ClassifiedBlock] = []
    all_spans = [s for page in extraction.pages for s in page.spans]
    if not all_spans:
        return classified
    body_size = statistics.median([s.size_pt for s in all_spans if s.size_pt])
    heading_threshold = body_size * 1.15

    for page in extraction.pages:
        columns = detect_columns(page.spans, page.page_width_pt)
        for column in columns:
            col_spans = _spans_in_column(page.spans, column)
            if not col_spans:
                continue
            lines = _group_lines(col_spans)
            block_lines_list = _group_blocks(lines)
            for block in block_lines_list:
                block_spans = [s for line in block for s in line]
                if not block_spans:
                    continue
                # Try v2's cross-row table detector first; fall back to v1 path.
                table_rows = _detect_table_v2(block)
                if table_rows is not None:
                    bbox = _line_bbox(block_spans)
                    classified.append(
                        ClassifiedBlock(
                            block_type="table",
                            text=None,
                            items=None,
                            rows=table_rows,
                            font=_majority_font(block_spans),
                            size_pt=round(_avg_size(block_spans), 1),
                            color_hex=_majority_color(block_spans),
                            bold=_is_bold(block_spans),
                            italic=sum(1 for s in block_spans if s.italic) > len(block_spans) / 2,
                            align=_line_align(bbox[0], bbox[2], page.page_width_pt),
                            bbox=bbox,
                            page_index=page.page_index,
                        )
                    )
                    continue
                cb = _classify_block_spans(
                    block, page.page_width_pt, page.page_index, body_size, heading_threshold
                )
                if cb is not None:
                    classified.append(cb)
    return classified


def classify_pdf_layout(extraction: PdfExtraction) -> list[ClassifiedBlock]:
    """Dispatch to v1 or v2 based on `VITEGRID_PARSER`. v1 stays the default."""
    parser_version = os.environ.get("VITEGRID_PARSER", "v1").lower()
    if parser_version == "v2":
        return _classify_pdf_layout_v2(extraction)
    return _classify_pdf_layout_v1(extraction)


# ---------------------------------------------------------------------------
# Headless Screenshot Engine
# ---------------------------------------------------------------------------


def render_layout_screenshot(layout_json_str: str, output_path: Path, width: int = 816, height: int = 1056):
    """
    Launches headless Chromium, loads the unbordered layout canvas route,
    and isolates the production container view to a raw PNG image.
    """
    import base64
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(
            viewport={"width": width, "height": height},
            device_scale_factor=1
        )
        page = context.new_page()

        # Pull dev endpoint parameters from settings strings
        frontend_url = os.environ.get("VITEGRID_FRONTEND_URL", "http://localhost:5173")
        page.goto(f"{frontend_url}/#/headless-preview")

        # Safely pass block metrics using standard base64 data transfers
        b64_layout = base64.b64encode(layout_json_str.encode("utf-8")).decode("utf-8")
        page.evaluate(f"localStorage.setItem('vitegrid_headless_layout', atob('{b64_layout}'))")

        page.reload()
        page.wait_for_selector("#headless-render-canvas", timeout=5000)

        canvas_element = page.query_selector("#headless-render-canvas")
        if canvas_element:
            canvas_element.screenshot(path=str(output_path))
        else:
            page.screenshot(path=str(output_path))

        browser.close()


# ---------------------------------------------------------------------------
# Visual Regression Math Engine
# ---------------------------------------------------------------------------


def calculate_visual_regression(ground_truth_path: Path, candidate_path: Path, diff_output_path: Path) -> float:
    """
    Evaluates pixel alignment variations via grayscale difference models,
    tints layout tracking anomalies in bright red, and returns total error percentages.
    """
    import cv2
    import numpy as np

    img_gt = cv2.imread(str(ground_truth_path), cv2.IMREAD_GRAYSCALE)
    img_cand = cv2.imread(str(candidate_path), cv2.IMREAD_GRAYSCALE)

    if img_gt is None or img_cand is None:
        raise ValueError("Could not read verification image sources.")

    if img_gt.shape != img_cand.shape:
        img_cand = cv2.resize(img_cand, (img_gt.shape[1], img_gt.shape[0]))

    diff = cv2.absdiff(img_gt, img_cand)
    _, thresh = cv2.threshold(diff, 15, 255, cv2.THRESH_BINARY)

    color_cand = cv2.imread(str(candidate_path))
    if color_cand.shape[:2] != img_gt.shape[:2]:
        color_cand = cv2.resize(color_cand, (img_gt.shape[1], img_gt.shape[0]))

    # Paint misalignment vectors explicitly in BGR Red [0, 0, 255]
    color_cand[thresh == 255] = [0, 0, 255]
    cv2.imwrite(str(diff_output_path), color_cand)

    mismatch_pixels = np.sum(thresh == 255)
    total_pixels = img_gt.shape[0] * img_gt.shape[1]
    return (mismatch_pixels / total_pixels) * 100

